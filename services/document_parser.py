"""
.docx text extraction (python-docx) + document type auto-detection.

Pure/local — no network. Runs in a threadpool from async callers.
"""
import io
from typing import Tuple

from docx import Document


class DocumentParseError(Exception):
    """Raised when a .docx file cannot be read."""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract full text from a .docx byte stream, preserving paragraph breaks and
    table cell content. Raises DocumentParseError on invalid files.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:  # python-docx raises a variety of errors on bad files
        raise DocumentParseError(str(exc)) from exc

    parts: list[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables (IDFs frequently use tables for Title/Inventor fields)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def detect_doc_type(raw_text: str) -> str:
    """
    Heuristic document-type detection.

    Returns 'claims', 'idf', or 'unknown'.
    """
    lowered = raw_text.lower()

    claims_signals = [
        "claim 1",
        "1. a computer-implemented method",
        "what is claimed is",
        "we claim",
        "i claim",
    ]
    idf_signals = [
        "invention disclosure",
        "idf",
        "description of invention",
        "invention disclosure form",
    ]

    has_claims = any(sig in lowered for sig in claims_signals)
    has_idf = any(sig in lowered for sig in idf_signals)

    # If both patterns appear, prefer the stronger/more specific signal.
    if has_idf and not has_claims:
        return "idf"
    if has_claims and not has_idf:
        return "claims"
    if has_claims and has_idf:
        # A claims doc rarely says "invention disclosure"; an IDF may quote a claim.
        # Prefer 'claims' only when an explicit numbered claim opener is present.
        if "1. a computer-implemented method" in lowered or "what is claimed is" in lowered:
            return "claims"
        return "idf"

    # Fall back on long-form prose w/ figure references => idf
    if "figure" in lowered or "fig." in lowered:
        return "idf"

    return "unknown"


def parse_and_detect(file_bytes: bytes) -> Tuple[str, str]:
    """Convenience: returns (raw_text, doc_type)."""
    raw_text = extract_text_from_docx(file_bytes)
    return raw_text, detect_doc_type(raw_text)
